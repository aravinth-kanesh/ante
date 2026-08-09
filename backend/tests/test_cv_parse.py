from io import BytesIO

from docx import Document

from app.services.cv_parse import extract_text


def test_extract_text_reads_docx_paragraphs_and_tables():
    doc = Document()
    doc.add_paragraph("Jane Doe")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, Linux, SQL"
    buffer = BytesIO()
    doc.save(buffer)

    text = extract_text("cv.docx", buffer.getvalue())
    assert "Jane Doe" in text  # a normal paragraph
    assert "Python, Linux, SQL" in text  # content that only lives in a table cell


def test_extract_text_plain_text():
    assert extract_text("cv.txt", b"Hello world") == "Hello world"


def test_extract_text_rejects_unsupported():
    import pytest

    with pytest.raises(ValueError):
        extract_text("cv.rtf", b"data")
