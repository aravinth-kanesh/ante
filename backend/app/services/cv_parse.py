import io
import re

from docx import Document
from pypdf import PdfReader

SUPPORTED = (".pdf", ".docx", ".txt")

# An upper bound on stored CV and job-description text. A real CV is a few thousand
# characters; this leaves generous headroom while keeping a pasted or extracted wall of
# text from flowing into the model prompt unbounded (a cost and context-window risk).
MAX_TEXT_CHARS = 50_000


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif name.endswith(".docx"):
        document = Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        # Many CVs lay their content out in tables, whose cell text is not in
        # `paragraphs`; include it so those CVs are not read as empty.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    parts.append(line)
        text = "\n".join(parts)
    elif name.endswith(".txt"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError("Unsupported file type; use PDF, Word (.docx) or plain text")

    # collapse runs of blank lines and trailing spaces
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_TEXT_CHARS]
